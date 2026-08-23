#!/usr/bin/env python3
"""
generate_word_report.py - Generates an Academic Word Document (.docx) Report
Creates a formatted academic report complete with tables, code snippets,
mathematical formulas, and embedded high-resolution figures.
"""

import os
import sys
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
REPORT_DIR = os.path.join(PROJECT_ROOT, "report")
FIGS_DIR = os.path.join(REPORT_DIR, "figures")
OUTPUT_DOCX = os.path.join(REPORT_DIR, "Dense_Matrix_Multiplication_Micro_Project_Report.docx")
ROOT_DOCX = os.path.join(PROJECT_ROOT, "Dense_Matrix_Multiplication_Micro_Project_Report.docx")

def set_cell_background(cell, color_hex):
    """Sets background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets cell padding."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def style_table(table, header_bg="1B365D", alt_bg="F4F6F9"):
    """Styles table with dark header, light alternating rows, and clean borders."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(table.rows):
        # Prevent row splitting across pages
        trPr = row._tr.get_or_add_trPr()
        trPr.append(OxmlElement('w:cantSplit'))
        
        if i == 0:
            trPr.append(OxmlElement('w:tblHeader'))
            for cell in row.cells:
                set_cell_background(cell, header_bg)
                set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.name = "Calibri"
                        run.font.size = Pt(10)
        else:
            bg = alt_bg if i % 2 == 1 else "FFFFFF"
            for cell in row.cells:
                set_cell_background(cell, bg)
                set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.name = "Calibri"
                        run.font.size = Pt(9.5)

def add_heading_styled(doc, text, level):
    """Adds headings with professional corporate/academic colors and spacing."""
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.runs[0]
    run.font.name = "Calibri"
    if level == 1:
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(27, 54, 93) # Navy
    elif level == 2:
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(41, 98, 153) # Steel Blue
    elif level == 3:
        run.font.size = Pt(11.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(60, 60, 60)
    return p

def add_callout(doc, text, title=""):
    """Adds shaded callout box for equations or notes."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "F0F4F8")
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    # Left border highlight
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="1B365D"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    if title:
        r_title = p.add_run(title + "\n")
        r_title.bold = True
        r_title.font.name = "Calibri"
        r_title.font.size = Pt(10.5)
        r_title.font.color.rgb = RGBColor(27, 54, 93)
    r = p.add_run(text)
    r.font.name = "Consolas" if "for (" in text or "C[" in text else "Calibri"
    r.font.size = Pt(9.5 if "for (" in text else 10.5)
    r.font.color.rgb = RGBColor(30, 30, 30)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_figure(doc, filename, caption, width=Inches(5.8)):
    """Inserts an image centered with a styled numbered caption."""
    img_path = os.path.join(FIGS_DIR, filename)
    if not os.path.exists(img_path):
        img_path = os.path.join(PROJECT_ROOT, "results", "graphs", filename)
    
    if os.path.exists(img_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(10)
        p_img.paragraph_format.space_after = Pt(3)
        p_img.paragraph_format.keep_with_next = True
        p_img.add_run().add_picture(img_path, width=width)
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_before = Pt(0)
        p_cap.paragraph_format.space_after = Pt(12)
        r = p_cap.add_run(caption)
        r.font.name = "Calibri"
        r.font.size = Pt(9.5)
        r.font.italic = True
        r.font.color.rgb = RGBColor(80, 80, 80)
    else:
        p = doc.add_paragraph(f"[Figure {filename} pending]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def build_report():
    doc = Document()
    
    # Page Setup: Standard Letter with 1 inch margins
    sections = doc.sections
    for s in sections:
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)

    # Base style
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(35, 35, 35)

    # ==========================================
    # COVER / TITLE BLOCK
    # ==========================================
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(24)
    p_title.paragraph_format.space_after = Pt(6)
    r_title = p_title.add_run("DENSE MATRIX MULTIPLICATION WITH CACHE BLOCKING")
    r_title.font.name = "Calibri"
    r_title.font.size = Pt(22)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(27, 54, 93)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(16)
    r_sub = p_sub.add_run("Tile Size Selection Tied to L1/L2 Cache Size and OpenMP collapse() Decisions")
    r_sub.font.name = "Calibri"
    r_sub.font.size = Pt(14)
    r_sub.font.bold = True
    r_sub.font.color.rgb = RGBColor(41, 98, 153)

    p_desc = doc.add_paragraph()
    p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_desc.paragraph_format.space_after = Pt(24)
    r_desc = p_desc.add_run("A Micro Project Report\nSubmitted in partial fulfillment of the requirements for the course in\nParallel and High Performance Computing")
    r_desc.font.name = "Calibri"
    r_desc.font.size = Pt(11)
    r_desc.font.italic = True
    r_desc.font.color.rgb = RGBColor(80, 80, 80)

    # Student metadata table
    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Student Name:", "Mahesh M"),
        ("USN / Roll Number:", "1MS21CS001"),
        ("Department:", "Computer Science and Engineering"),
        ("Institution:", "Department of CSE / High Performance Computing Lab"),
        ("Academic Year:", "2026–2027")
    ]
    for row_idx, (k, v) in enumerate(meta_data):
        cell_k = meta_table.cell(row_idx, 0)
        cell_v = meta_table.cell(row_idx, 1)
        cell_k.width = Inches(2.2)
        cell_v.width = Inches(4.3)
        set_cell_background(cell_k, "F4F6F9")
        set_cell_background(cell_v, "FFFFFF")
        set_cell_margins(cell_k, top=60, bottom=60, left=100, right=100)
        set_cell_margins(cell_v, top=60, bottom=60, left=100, right=100)
        
        pk = cell_k.paragraphs[0]
        pv = cell_v.paragraphs[0]
        rk = pk.add_run(k)
        rk.font.bold = True
        rk.font.size = Pt(10)
        rv = pv.add_run(v)
        rv.font.size = Pt(10)

    doc.add_page_break()

    # ==========================================
    # ABSTRACT
    # ==========================================
    add_heading_styled(doc, "ABSTRACT", level=1)
    p_abs = doc.add_paragraph()
    p_abs.paragraph_format.line_spacing = 1.15
    p_abs.paragraph_format.space_after = Pt(10)
    p_abs.add_run(
        "Dense matrix multiplication is a fundamental computational operation used in scientific computing, numerical linear algebra, "
        "computer vision, physical simulations, and machine-learning workloads. Although the mathematical formulation is straightforward, "
        "its performance is strongly influenced by memory-access patterns, cache hierarchy, compiler optimization, and available processor parallelism.\n\n"
        "This project investigates the performance optimization of dense square matrix multiplication using loop reordering, cache blocking, and OpenMP-based parallelization. "
        "The study specifically examines the relationship between tile size and the L1/L2 cache hierarchy and evaluates the effect of OpenMP collapse(2) on workload distribution.\n\n"
        "Five principal implementation strategies were investigated: conventional ijk matrix multiplication, loop-reordered ikj multiplication, serial cache-blocked multiplication, "
        "OpenMP-parallel multiplication, and OpenMP cache-blocked multiplication with and without collapse(2). A theoretical three-tile working-set model was used to identify "
        "cache-compatible tile-size candidates. The experimental evaluation was conducted on an Intel Core Ultra 5 225H processor with 48 KB per-core L1D cache, 2 MB per-core/cluster L2 cache, "
        "and an 18 MB shared L3 cache.\n\n"
        "The implementation was validated using 81 automated correctness tests, with all 81 tests passing and a maximum reported absolute difference of 0.00e+00. Experimental measurements "
        "demonstrated a substantial improvement from loop reordering, while cache blocking and OpenMP parallelization provided further performance gains. The highest reported throughput in the "
        "investigated configuration was 52.376 GFLOPS at N=1024 (and 66.49 GFLOPS at N=1536) using the blocked OpenMP implementation with collapse(2), achieving up to a 91.10x speedup relative to the reported naive serial baseline.\n\n"
        "The results demonstrate that theoretical cache capacity is useful for narrowing the tile-size search space, but the experimentally optimal tile size cannot be determined from cache capacity alone. "
        "OpenMP collapse(2) also provides additional parallel iteration space when the number of tile rows is insufficient to fully utilize the available processor threads."
    )

    # ==========================================
    # 1. INTRODUCTION
    # ==========================================
    add_heading_styled(doc, "1. INTRODUCTION", level=1)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "Dense matrix multiplication computes the product of two matrices according to C = A x B. "
        "For two square matrices of dimension N x N, each element of the output matrix is calculated as:\n\n"
    )
    add_callout(doc, "C[i, j] = Σ (A[i, k] * B[k, j])   for k = 0 to N-1", "Matrix Multiplication Definition")
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "The operation requires approximately 2N³ floating-point operations (N³ multiplications and N³ additions). "
        "Matrix multiplication is therefore computationally intensive for large matrix dimensions. At the same time, performance depends not only on the number of arithmetic operations "
        "but also on how efficiently the processor accesses data from the memory hierarchy.\n\n"
        "Modern processors use multiple levels of cache between the CPU execution units and main memory. L1 cache provides very low latency (~4-5 cycles) but limited capacity, "
        "while L2 and L3 caches provide progressively larger storage at higher access latency. Consequently, an implementation that repeatedly accesses data in an unfavorable pattern "
        "can experience significant performance degradation ('the Memory Wall') even when the underlying mathematical algorithm is correct.\n\n"
        "This project investigates several optimization techniques designed to improve both memory locality and processor utilization. The main techniques studied are loop reordering, "
        "cache blocking, OpenMP parallelization, and OpenMP iteration-space collapsing."
    )

    # ==========================================
    # 2. PROBLEM STATEMENT
    # ==========================================
    add_heading_styled(doc, "2. PROBLEM STATEMENT", level=1)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "The conventional ijk matrix multiplication algorithm accesses matrix B column-wise while the matrix is stored in row-major order. The inner operation is:\n\n"
    )
    add_callout(doc, "C[i][j] += A[i][k] * B[k][j];", "Inner Loop Operation (ijk)")
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "As k changes, the address of B[k][j] changes by approximately N elements (N * 8 bytes). For large matrices, this creates a strided memory-access pattern and degrades spatial locality. "
        "A 64-byte cache line can contain eight double-precision values (64 / 8 = 8). When the computation accesses only one value from a fetched cache line before jumping to another distant location, "
        "7 out of 8 values in the transferred cache line are evicted before reuse, resulting in severe bandwidth waste and CPU pipeline stalls.\n\n"
        "The fundamental research problem addressed in this project is:"
    )
    add_callout(doc, "How can dense matrix multiplication be organized so that the processor obtains better cache locality while simultaneously exploiting multicore parallelism, and how does OpenMP collapse(2) influence thread scalability and cache locality?", "Core Research Question")

    # ==========================================
    # 3. OBJECTIVES
    # ==========================================
    add_heading_styled(doc, "3. OBJECTIVES", level=1)
    objectives = [
        "Implement conventional dense matrix multiplication using the ijk loop ordering.",
        "Implement the cache-friendly ikj loop ordering with unit-stride inner memory access.",
        "Implement 6-loop cache-blocked matrix multiplication with dynamic boundary clamping.",
        "Analyze the mathematical relationship between tile size B and CPU L1/L2 cache capacity.",
        "Implement OpenMP parallel matrix multiplication.",
        "Implement OpenMP cache-blocked matrix multiplication.",
        "Investigate OpenMP collapse(2) concurrency, work distribution, and false sharing.",
        "Measure execution time, standard deviation, and computational throughput (GFLOPS).",
        "Calculate speedup relative to the naive serial baseline.",
        "Verify numerical correctness of all optimized implementations against serial reference.",
        "Compare theoretical cache-based tile selection with experimentally measured optimums.",
        "Generate reproducible empirical performance datasets and publication-ready charts."
    ]
    for idx, obj in enumerate(objectives, 1):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(3)
        r_num = p.add_run(f"{idx}. ")
        r_num.bold = True
        r_num.font.color.rgb = RGBColor(27, 54, 93)
        p.add_run(obj)

    # ==========================================
    # 4. SYSTEM ENVIRONMENT
    # ==========================================
    add_heading_styled(doc, "4. SYSTEM ENVIRONMENT", level=1)
    p = doc.add_paragraph("The experiments were performed on the physical hardware and software environment detailed below:")
    p.paragraph_format.space_after = Pt(6)

    env_table = doc.add_table(rows=16, cols=2)
    env_rows = [
        ("Component / Parameter", "Detected Hardware / Software Specification"),
        ("Operating System", "Windows 11 AMD64 (Build 26200)"),
        ("Processor (CPU)", "Intel Core Ultra 5 225H"),
        ("Physical Cores", "14 Cores (4 P-cores + 8 E-cores + 2 LP E-cores)"),
        ("Logical Processors", "14 Execution Threads"),
        ("L1 Data Cache (L1D)", "48 KB per Performance Core (32 KB per Efficient Core)"),
        ("L1 Instruction Cache (L1I)", "64 KB per Core"),
        ("Aggregate L1 Cache", "1,408 KB (1.4 MB total across cores)"),
        ("L2 Unified Cache", "2,048 KB (2.0 MB) per P-core / E-cluster"),
        ("Aggregate L2 Cache", "22,528 KB (22.0 MB total)"),
        ("L3 Last-Level Cache", "18,432 KB (18.0 MB shared LLC)"),
        ("Cache Line Size", "64 bytes (8 double-precision floats)"),
        ("C Compiler", "GCC 16.2.0 (MinGW-W64 UCRT with OpenMP)"),
        ("Compiler Flags", "-O3 -Wall -Wextra -std=c11 -fopenmp -march=native"),
        ("Python Environment", "Python 3.11.9 (64-bit)"),
        ("Analysis Libraries", "NumPy 2.4.6, Pandas 3.0.5, Matplotlib 3.11.1, psutil 7.2.2")
    ]
    for r_idx, (c1, c2) in enumerate(env_rows):
        env_table.cell(r_idx, 0).paragraphs[0].text = c1
        env_table.cell(r_idx, 1).paragraphs[0].text = c2
        env_table.cell(r_idx, 0).width = Inches(2.5)
        env_table.cell(r_idx, 1).width = Inches(4.0)
    style_table(env_table)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ==========================================
    # 5. THEORETICAL BACKGROUND
    # ==========================================
    add_heading_styled(doc, "5. THEORETICAL BACKGROUND", level=1)
    
    add_heading_styled(doc, "5.1 Matrix Multiplication", level=2)
    p = doc.add_paragraph("For dense square matrices A, B, and C in R^(NxN), total arithmetic operations equal exactly 2N³ FLOPs.")
    p.paragraph_format.line_spacing = 1.15

    add_heading_styled(doc, "5.2 Memory Locality", level=2)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.add_run("Two fundamental principles govern cache efficiency in high-performance computing:\n")
    p.add_run("• Spatial Locality: ").bold = True
    p.add_run("When a memory address is accessed, contiguous adjacent addresses within the same 64-byte cache line are loaded automatically.\n")
    p.add_run("• Temporal Locality: ").bold = True
    p.add_run("Data loaded into low-latency L1/L2 cache should be repeatedly reused across multiple arithmetic operations before being evicted to main memory.")

    add_heading_styled(doc, "5.3 Loop Reordering (ijk vs ikj)", level=2)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "In the conventional ijk loop arrangement, matrix B is traversed with stride N across k, causing high cache miss rates. "
        "In the optimized ikj arrangement, the inner j loop iterates across adjacent columns, converting both B and C access into contiguous unit-stride row streaming. "
        "Furthermore, A[i*N + k] remains invariant across the inner j loop and is held inside a CPU register throughout the vector computation."
    )

    # ==========================================
    # 6. CACHE BLOCKING
    # ==========================================
    add_heading_styled(doc, "6. CACHE BLOCKING (TILING)", level=1)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "Cache blocking decomposes the N x N matrix into sub-matrices of dimension B x B (tiles). "
        "By restricting inner computation to sub-blocks, the working set of tiles remains resident in fast per-core L1D/L2 caches.\n\n"
        "The 6-loop blocked structure is implemented with dynamic boundary handling:"
    )
    add_callout(doc, 
        "for (int ii = 0; ii < N; ii += B) {\n"
        "    int i_max = MIN(ii + B, N);\n"
        "    for (int kk = 0; kk < N; kk += B) {\n"
        "        int k_max = MIN(kk + B, N);\n"
        "        for (int jj = 0; jj < N; jj += B) {\n"
        "            int j_max = MIN(jj + B, N);\n"
        "            for (int i = ii; i < i_max; i++) {\n"
        "                for (int k = kk; k < k_max; k++) {\n"
        "                    double a_ik = A[i * N + k];\n"
        "                    for (int j = jj; j < j_max; j++) {\n"
        "                        C[i * N + j] += a_ik * B[k * N + j];\n"
        "                    }\n"
        "                }\n"
        "            }\n"
        "        }\n"
        "    }\n"
        "}", "Cache-Blocked ikj Matrix Multiplication Kernel"
    )

    # ==========================================
    # 7. CACHE-AWARE TILE-SIZE MODEL
    # ==========================================
    add_heading_styled(doc, "7. CACHE-AWARE TILE-SIZE MODEL", level=1)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "For double-precision floating point data (8 bytes/element), the memory footprint of a single square tile is W_tile = B² * 8 bytes. "
        "A simplified three-tile working-set model accounts for concurrent residency of an active A-tile, B-tile, and C-tile:\n\n"
    )
    add_callout(doc, "W_3-tile ≈ Footprint(A_tile) + Footprint(B_tile) + Footprint(C_tile) = 3 * B² * 8 bytes = 24 * B² bytes", "Three-Tile Working Set Formulation")

    add_heading_styled(doc, "7.1 L1D & L2 Cache Capacity Bounds", level=2)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "• L1D Cache Bound (48 KB): Solving 24 * B² ≤ 48 * 1024 yields B ≤ 45.2. Thus, tile sizes B ∈ {8, 16, 32} fit strictly within L1D cache.\n"
        "• L2 Cache Bound (2 MB): Solving 24 * B² ≤ 2 * 1024² yields B ≤ 295.4. Thus, tile sizes B ∈ {48, 64, 96, 128, 192, 256} fit within the 2.0 MB L2 cache."
    )

    # Working set table
    ws_table = doc.add_table(rows=10, cols=6)
    ws_data = [
        ("Tile Size (B)", "Single Tile", "3-Tile Working Set", "% L1D (48 KB)", "% L2 (2 MB)", "Theoretical Classification"),
        ("8 x 8", "0.50 KB", "1.50 KB", "3.1 %", "0.07 %", "L1D Cache Bound"),
        ("16 x 16", "2.00 KB", "6.00 KB", "12.5 %", "0.29 %", "L1D Cache Bound"),
        ("32 x 32", "8.00 KB", "24.00 KB", "50.0 %", "1.17 %", "L1D Cache Bound (Optimal L1 candidate)"),
        ("48 x 48", "18.00 KB", "54.00 KB", "112.5 %", "2.64 %", "L2 Cache Bound"),
        ("64 x 64", "32.00 KB", "96.00 KB", "200.0 %", "4.69 %", "L2 Cache Bound (Optimal L2 candidate)"),
        ("96 x 96", "72.00 KB", "216.00 KB", "450.0 %", "10.55 %", "L2 Cache Bound"),
        ("128 x 128", "128.00 KB", "384.00 KB", "800.0 %", "18.75 %", "L2 Cache Bound (Peak SIMD candidate)"),
        ("192 x 192", "288.00 KB", "864.00 KB", "1800.0 %", "42.19 %", "L2 Cache Bound"),
        ("256 x 256", "512.00 KB", "1536.00 KB", "3200.0 %", "75.00 %", "Near L2 Cache Limit")
    ]
    for r_idx, row in enumerate(ws_data):
        for c_idx, val in enumerate(row):
            ws_table.cell(r_idx, c_idx).paragraphs[0].text = val
    style_table(ws_table)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    add_figure(doc, "07_cache_tile_analysis.png", "Figure 1: Mathematical L1D and L2 Cache Capacity Bounds vs. Empirical Throughput (GFLOPS).")

    # ==========================================
    # 8. OPENMP PARALLELIZATION & 9. COLLAPSE(2)
    # ==========================================
    add_heading_styled(doc, "8. OPENMP PARALLELIZATION & 9. COLLAPSE(2)", level=1)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "OpenMP is used to distribute matrix computation across the 14 available CPU cores. "
        "In a conventional parallel blocked implementation, only the outermost ii loop is parallelized (#pragma omp parallel for). "
        "For matrix size N=1024 and tile size B=128, the number of outer tile strips is:\n\n"
    )
    add_callout(doc, "Outer Tile Iterations = ⌈1024 / 128⌉ = 8 work items", "Single-Loop Parallel Iteration Space")
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "Because only 8 iterations exist for 14 physical cores, 6 cores remain idle (thread starvation). "
        "To resolve this, OpenMP collapse(2) flattens the 2D grid of tile loops (ii, jj) into a unified 2D iteration space:\n\n"
    )
    add_callout(doc, "Collapsed Iterations = ⌈1024 / 128⌉ * ⌈1024 / 128⌉ = 8 * 8 = 64 work items", "Collapsed 2D Iteration Space (collapse(2))")
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "With 64 independent work units, all 14 cores receive balanced chunks (~4.5 tiles/thread). "
        "Because each iteration of (ii, jj) writes exclusively to a disjoint sub-matrix C[ii..i_max, jj..j_max], there are zero write data races."
    )

    add_figure(doc, "05_collapse_comparison.png", "Figure 2: Execution Time Comparison between OpenMP Blocked (No Collapse) and collapse(2) across Thread Counts.")

    # ==========================================
    # 10. EXPERIMENTAL METHODOLOGY & 11. CORRECTNESS
    # ==========================================
    add_heading_styled(doc, "10. EXPERIMENTAL METHODOLOGY & 11. CORRECTNESS", level=1)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "To guarantee high academic integrity and reproducibility, the benchmark engine incorporates strict experimental controls:\n"
        "• 64-Byte Memory Alignment: All matrix buffers are aligned to cache-line boundaries via _aligned_malloc / posix_memalign.\n"
        "• Deterministic Initialization: Matrices are generated with bounded sinusoidal values to prevent floating point overflow.\n"
        "• Monotonic High-Resolution Timing: Timed with omp_get_wtime() after untimed warmups with 5 repetitions per configuration.\n"
        "• Unit Test Suite: 81 automated tests evaluated matrix dimensions 1x1 to 257x257 with an absolute error threshold of 1e-9."
    )

    # Correctness table
    corr_table = doc.add_table(rows=5, cols=2)
    corr_rows = [
        ("Metric / Validation Parameter", "Observed Result"),
        ("Total Automated Test Cases", "81 Tests"),
        ("Passed Tests", "81 (100.0% Success Rate)"),
        ("Failed Tests", "0"),
        ("Maximum Observed Absolute Difference", "0.00e+00 (Numerical Exactness)")
    ]
    for r_idx, (c1, c2) in enumerate(corr_rows):
        corr_table.cell(r_idx, 0).paragraphs[0].text = c1
        corr_table.cell(r_idx, 1).paragraphs[0].text = c2
        corr_table.cell(r_idx, 0).width = Inches(3.2)
        corr_table.cell(r_idx, 1).width = Inches(3.3)
    style_table(corr_table)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ==========================================
    # 12. PERFORMANCE RESULTS
    # ==========================================
    add_heading_styled(doc, "12. PERFORMANCE RESULTS", level=1)
    p = doc.add_paragraph("Table 1 summarizes the empirical performance measurements for N=1024 across all kernel variants:")
    p.paragraph_format.space_after = Pt(6)

    res_table = doc.add_table(rows=10, cols=6)
    res_data = [
        ("Kernel Variant", "Tile (B)", "Threads", "Median Time (s)", "Throughput (GFLOPS)", "Speedup vs Naive"),
        ("Naive ijk", "—", "1", "3.7348 s", "0.575 GF", "1.00x"),
        ("IKJ (Loop Reordered)", "—", "1", "0.3441 s", "6.241 GF", "10.85x"),
        ("Blocked Serial", "16", "1", "0.4789 s", "4.484 GF", "7.80x"),
        ("Blocked Serial", "32", "1", "0.4190 s", "5.125 GF", "8.91x"),
        ("Blocked Serial", "64", "1", "0.3662 s", "5.864 GF", "10.20x"),
        ("Blocked Serial", "128", "1", "0.3392 s", "6.331 GF", "11.01x"),
        ("OpenMP IKJ", "—", "14", "0.0519 s", "41.376 GF", "71.96x"),
        ("OpenMP Blocked", "64", "14", "0.0460 s", "46.683 GF", "81.19x"),
        ("OpenMP Blocked + collapse(2)", "64", "14", "0.0410 s", "52.376 GF", "91.10x")
    ]
    for r_idx, row in enumerate(res_data):
        for c_idx, val in enumerate(row):
            res_table.cell(r_idx, c_idx).paragraphs[0].text = val
    style_table(res_table)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    add_figure(doc, "03_serial_vs_openmp.png", "Figure 3: Execution Time and GFLOPS Comparison across Kernel Implementations (N=1024).")
    add_figure(doc, "01_execution_time_tile_size.png", "Figure 4: Execution Time vs. Cache Blocking Tile Size B (N=1024).")
    add_figure(doc, "02_gflops_tile_size.png", "Figure 5: Computational Throughput (GFLOPS) vs. Tile Size B.")
    add_figure(doc, "04_thread_scaling.png", "Figure 6: Parallel Speedup vs. OpenMP Thread Count (1 to 14 Threads).")
    add_figure(doc, "06_matrix_size_scaling.png", "Figure 7: Performance Scaling across Matrix Dimensions (N=512, 1024, 1536).")
    add_figure(doc, "08_tile_thread_heatmap.png", "Figure 8: 2D Performance Heatmap (Tile Size vs. Thread Count in GFLOPS).")

    # ==========================================
    # 13 - 17. IN-DEPTH ANALYSIS
    # ==========================================
    add_heading_styled(doc, "13. ANALYSIS OF LOOP REORDERING", level=1)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "Transitioning from ijk to ikj produced an immediate 10.85x speedup in single-threaded execution (reducing runtime from 3.7348s to 0.3441s). "
        "This confirms that memory layout alignment and spatial locality dominate performance in dense matrix algorithms. "
        "In the ikj loop, contiguous row streaming allows hardware prefetchers and AVX2 vector pipelines to operate at maximum efficiency."
    )

    add_heading_styled(doc, "14. ANALYSIS OF TILE SIZE SELECTION", level=1)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "While theoretical modeling identified B=32 as the candidate matching the 48 KB L1D cache, experimental measurements demonstrated that B=64 and B=128 "
        "yield the highest computational throughput (reaching 6.331 GFLOPS in serial and 52.376 GFLOPS in parallel). "
        "This occurs because larger tiles amortize loop branch overheads and permit deep compiler vector unrolling (FMA instructions) while staying well within the 2.0 MB L2 cache."
    )

    add_heading_styled(doc, "15. OPENMP SCALING & 16. COLLAPSE(2) ANALYSIS", level=1)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "OpenMP multithreading scaled throughput from 6.24 GFLOPS to 52.38 GFLOPS across 14 cores. "
        "The collapse(2) implementation outperformed non-collapsed blocked OpenMP (52.38 GF vs 46.68 GF) by expanding the iteration space from 16 work items to 256 work items, "
        "ensuring uniform load distribution across all 14 execution threads."
    )

    add_heading_styled(doc, "17. OVERALL PEAK CONFIGURATION", level=1)
    best_table = doc.add_table(rows=9, cols=2)
    best_rows = [
        ("Parameter", "Measured Optimal Value"),
        ("Matrix Dimension (N)", "1024 x 1024 (and 1536 x 1536)"),
        ("Optimal Tile Size (B)", "64 x 64 (and 128 x 128)"),
        ("Optimal Thread Count", "14 OpenMP Threads"),
        ("Optimal Kernel Variant", "OpenMP Blocked + collapse(2)"),
        ("Median Execution Time", "0.0410 s (N=1024) / 0.1090 s (N=1536)"),
        ("Peak Computational Throughput", "52.376 GFLOPS (N=1024) / 66.49 GFLOPS (N=1536)"),
        ("Speedup vs Naive Baseline", "91.10x Speedup"),
        ("Speedup vs Serial ikj", "8.39x Parallel Speedup")
    ]
    for r_idx, (c1, c2) in enumerate(best_rows):
        best_table.cell(r_idx, 0).paragraphs[0].text = c1
        best_table.cell(r_idx, 1).paragraphs[0].text = c2
        best_table.cell(r_idx, 0).width = Inches(2.8)
        best_table.cell(r_idx, 1).width = Inches(3.7)
    style_table(best_table)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ==========================================
    # 18 - 21. DISCUSSION, LIMITATIONS & CONCLUSION
    # ==========================================
    add_heading_styled(doc, "18. ADVANTAGES OF PROPOSED APPROACH", level=1)
    advs = [
        "Superior Memory Locality: Unit-stride traversal maximizes 64-byte cache line utilization.",
        "Effective Cache Hierarchy Residency: Tiling keeps active sub-matrices within per-core L1D/L2 caches.",
        "Scalable OpenMP Parallelism: Efficient multicore scaling up to 14 threads.",
        "Enhanced Scheduling via collapse(2): Eliminates thread under-utilization for large tiles.",
        "100% Correctness Validation: Zero numerical discrepancy across 81 unit test configurations."
    ]
    for a in advs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.add_run("• ").bold = True
        p.add_run(a)

    add_heading_styled(doc, "19. LIMITATIONS & 20. FUTURE SCOPE", level=1)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "Limitations include simplified three-tile cache assumptions (which omit set-associativity conflict misses) and core frequency scaling on hybrid architectures. "
        "Future scope includes integrating BLIS-style multi-level packing, explicit SIMD intrinsics (AVX-512), NUMA first-touch policies, and GPU offloading via OpenMP target directives."
    )

    add_heading_styled(doc, "21. CONCLUSION", level=1)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "This project comprehensively evaluated dense matrix multiplication optimization across memory hierarchy and parallel paradigms. "
        "Memory-access order is the single most critical factor in dense computing: swapping loops from ijk to ikj delivered a 10.85x serial speedup. "
        "Cache blocking controlled the active working set, with theoretical models successfully narrowing candidate search spaces to L1D (B<=32) and L2 (B<=256) bounds. "
        "OpenMP collapse(2) maximized parallel efficiency across 14 cores, achieving a peak throughput of 52.376 GFLOPS at N=1024 (and 66.49 GFLOPS at N=1536), "
        "representing a 91.10x speedup over the naive serial implementation with 100% numerical correctness."
    )

    add_heading_styled(doc, "REFERENCES", level=1)
    refs = [
        "OpenMP Architecture Review Board, OpenMP Application Programming Interface Specification, Version 5.2.",
        "Intel Corporation, Intel 64 and IA-32 Architectures Software Developer's Manual: Volume 3B (System Programming Guide).",
        "Free Software Foundation, GNU Compiler Collection (GCC) Documentation - Optimization and OpenMP Directives.",
        "Hennessy, J. L., & Patterson, D. A., Computer Architecture: A Quantitative Approach, 6th Edition, Morgan Kaufmann.",
        "Project experimental dataset and methodology documentation: methodology.md, cache_analysis.md, and experiment_notes.md."
    ]
    for idx, r in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.add_run(f"[{idx}] ").bold = True
        p.add_run(r)

    # Save to both report/ and project root
    doc.save(OUTPUT_DOCX)
    doc.save(ROOT_DOCX)
    print(f"[+] Word report successfully generated:\n    -> {OUTPUT_DOCX}\n    -> {ROOT_DOCX}")

if __name__ == "__main__":
    build_report()
